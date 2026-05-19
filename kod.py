# ==============================================================================
# 📚 1. KÜTÜPHANELER VE GÜVENLİK AYARLARI
# ==============================================================================
import ssl
ssl._create_default_https_context = ssl._create_unverified_context

import os       
import base64   
import easyocr  
import json     
from PIL import Image, ImageEnhance 
import google.generativeai as genai 
import streamlit as st  
import pandas as pd        
from docx import Document  
from fpdf import FPDF  
from io import BytesIO     

if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
else:
    st.error("Lütfen Streamlit Secrets alanına GEMINI_API_KEY anahtarınızı ekleyin.")

# 🖥️ SAYFA AYARLARI
st.set_page_config(page_title="PalaeoLab AI - Evrensel Arşiv ve Analiz Sistemi", layout="wide")

# ==============================================================================
# 🎨 2. PREMIUM GÖRSEL TASARIM VE KUSURSUZ GİZLEME AYARLARI (CSS KATMANI)
# ==============================================================================
banner_adi = "banner.png"
bg_image_html = ""

if os.path.exists(banner_adi):
    with open(banner_adi, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode()
    bg_image_html = f"""
    <style>
    .stApp {{
        background-image: url("data:image/png;base64,{encoded_string}") !important;
        background-size: 100% 140px !important; 
        background-position: top center !important;
        background-repeat: no-repeat !important;
        background-color: #080c10 !important; 
    }}
    </style>
    """

st.markdown(f"""
    {bg_image_html}
    <style>
    @import url('https://googleapis.com');
    
    .stApp {{ 
        font-family: 'Inter', sans-serif !important;
        color: #f1f5f9 !important; 
    }}
    
    h1, h2, h3, p, label {{ font-family: 'Inter', sans-serif !important; color: #ffffff !important; }}
    .main .block-container {{ padding-top: 160px !important; }}
    
    header[data-testid="stHeader"] button, header[data-testid="stHeader"] div {{ 
        color: #34d399 !important; 
        font-weight: bold !important; 
    }}
    header[data-testid="stHeader"] svg {{ fill: #34d399 !important; }}
    
    /* 🛠️ SAĞ ALTAKİ TÜM REKLAM VE YÖNETİM BUTONLARINI TAMAMEN YOK EDEN SİHİRLİ KOD */
    footer, div[data-testid="stDecoration"], .viewerBadge_container__16vsn, div[class*="manageApp"], div[class*="viewerBadge"] {{ 
        visibility: hidden !important; 
        display: none !important; 
    }}
    
    .adim-karti {{ 
        background: rgba(13, 20, 32, 0.85) !important; 
        backdrop-filter: blur(16px) !important;
        -webkit-backdrop-filter: blur(16px) !important;
        padding: 24px !important; 
        border-radius: 16px !important; 
        border: 1px solid rgba(52, 211, 153, 0.25) !important; 
        margin-top: 25px !important; 
        margin-bottom: 25px !important;
        box-shadow: 0 12px 40px 0 rgba(0, 0, 0, 0.5) !important;
    }}
    
    div[data-testid="stFileUploader"] {{ 
        border: 2px dashed rgba(52, 211, 153, 0.4) !important; 
        border-radius: 14px !important; 
        background-color: rgba(9, 14, 23, 0.7) !important; 
        padding: 25px !important; 
    }}
    div[data-testid="stFileUploader"] * {{ color: #ffffff !important; }}
    
    div[data-testid="stRadio"] p, div[data-testid="stSlider"] p, label p, .stWidgetLabel, div[data-testid="stRadio"] label {{
        color: #ffffff !important; 
        font-weight: 600 !important;
    }}
    
    div[data-testid="stSlider"] [data-handle="true"] {{
        background-color: #10b981 !important;
        box-shadow: 0px 0px 10px #10b981 !important;
    }}
    div[data-testid="stSlider"] [data-baseweb="slider"] > div {{
        background: linear-gradient(to right, #10b981, #059669) !important;
    }}
    
    .stButton>button[kind="primary"] {{ 
        background: linear-gradient(135deg, #10b981, #059669) !important; 
        color: #ffffff !important; 
        border-radius: 10px !important; 
        border: none !important;
        font-weight: 600 !important; 
        font-size: 16px !important; 
        padding: 14px 20px !important;
        box-shadow: 0 4px 14px 0 rgba(16, 185, 129, 0.3) !important;
        width: 100% !important;
    }}
    
    .stTextArea textarea {{
        font-family: 'JetBrains Mono', monospace !important;
        font-size: 14px !important;
        background-color: rgba(7, 11, 19, 0.9) !important;
        color: #a7f3d0 !important;
        border: 1px solid rgba(52, 211, 153, 0.3) !important;
        border-radius: 12px !important;
    }}
    
    .ana-baslik {{ font-size: 38px !important; font-weight: 800 !important; color: #ffffff !important; text-align: center; margin-bottom: 5px; }}
    .ana-baslik span {{ background: linear-gradient(to right, #34d399, #10b981); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }}
    .alt-baslik {{ text-align: center !important; color: #94a3b8 !important; font-size: 16px !important; margin-bottom: 35px; }}
    
    .stImage img {{ border-radius: 12px !important; border: 1px solid rgba(255, 255, 255, 0.1) !important; }}
    
    section[data-testid="stSidebar"] {{
        background-color: #070b13 !important;
        border-right: 1px solid rgba(52, 211, 153, 0.15) !important;
    }}
    button[data-testid="stSidebarCollapseButton"] {{
        color: #34d399 !important;
        background-color: rgba(20, 27, 38, 0.9) !important;
        border-radius: 8px !important;
        border: 1px solid rgba(52, 211, 153, 0.3) !important;
        margin-top: 145px !important; 
    }}
    
    .sozluk-kart {{
        background: rgba(16, 185, 129, 0.1) !important;
        border-left: 4px solid #10b981 !important;
        padding: 10px 15px !important;
        border-radius: 4px 8px 8px 4px !important;
        margin-bottom: 10px !important;
    }}
    </style>
""", unsafe_allow_html=True)

st.markdown('<p class="ana-baslik">🔬 PALAEO<span>LAB</span> AI</p>', unsafe_allow_html=True)
st.markdown('<p class="alt-baslik">Osmanlıca Paleografi ve Tarihi Arşivler İçin Yeni Nesil Yapay Zekâ Platformu</p>', unsafe_allow_html=True)

# ==============================================================================
# 🗄️ 3. GELİŞMİŞ OTURUM HAFIZASI (SESSION STATE)
# ==============================================================================
if "belge_arsivi" not in st.session_state: 
    st.session_state.belge_arsivi = {}

if "aktif_belge_adi" not in st.session_state: 
    st.session_state.aktif_belge_adi = None

# ==============================================================================
# 📁 4. GIZLI SOL MENÜ (SIDEBAR) GEÇMİŞ PANELİ YÖNETİMİ
# ==============================================================================
with st.sidebar:
    st.markdown("### 🗄️ Laboratuvar Arşivi")
    st.write("Oturum geçmişinizdeki belgelere buradan ulaşabilirsiniz. Sol üstteki üç çizgili butona (☰) basarak bu paneli tamamen gizleyebilirsiniz.")
    st.write("---")
    
    if st.session_state.belge_arsivi:
        arsiv_listesi = list(st.session_state.belge_arsivi.keys())
        secilen_belge = st.radio("Geçmiş Belgeler", arsiv_listesi, label_visibility="collapsed")
        st.session_state.aktif_belge_adi = secilen_belge
    else:
        st.info("Henüz taranmış bir geçmiş döküman bulunmuyor.")
        
    st.write("---")
    if st.button("🗑️ Geçmişi Temizle"):
        st.session_state.belge_arsivi = {}
        st.session_state.aktif_belge_adi = None
        st.rerun()

# ==============================================================================
# 🎛️ 5. ADIM 1: GÖRSEL İYİLEŞTİRME VE DOSYA YÜKLEME PANELİ
# ==============================================================================
st.markdown("""
<div class="adim-karti">
    ⚡ <b>ADIM 1: Belge İyileştirme ve Yükleme Paneli</b><br>
    Yapay zekâ analizini başlatmadan önce belgenin netliğini ve kontrastını optimize edin.
</div>
""", unsafe_allow_html=True)

col_ayar, col_yukle = st.columns(2)

with col_ayar:
    st.write("⚙️ **Görsel Ön İşleme Katmanı**")
    filtre_modu = st.radio("Renk Modu", ["Orijinal Renk Spektrumu", "Siyah-Beyaz Stüdyo (Önerilen)"], horizontal=True)
    kontrast_seviyesi = st.slider("Metin Belirginleştirme (Kontrast)", 1.0, 3.0, 1.6, step=0.1)

with col_yukle:
    yuklenen_dosya = st.file_uploader("Tarihi belgenizi buraya sürükleyip bırakın...", type=["jpg", "jpeg", "png", "jfif"])

# ==============================================================================
# 🔮 6. HİBRİT OKUMA VE ANALİZ MOTORU
# ==============================================================================
if yuklenen_dosya is not None:
    orijinal_resim = Image.open(yuklenen_dosya)
    resim_islem = orijinal_resim.convert('RGB') if orijinal_resim.mode in ('RGBA', 'LA') else orijinal_resim
    
    if filtre_modu == "Siyah-Beyaz Stüdyo (Önerilen)":
        resim_islem = resim_islem.convert('L').convert('RGB')
        
    gelistirici = ImageEnhance.Contrast(resim_islem)
    resim_islem = gelistirici.enhance(kontrast_seviyesi)

    st.write("")
    if st.button("🔮 Akıllı Çözümleme ve Analizi Başlat", type="primary"):
        with st.spinner("🧠 Belge katmanları inceleniyor ve paleografik yazılar çözülüyor..."):
            try:
                gecici_yol = "gecici_resim.jpg"
                resim_islem.save(gecici_yol, format="JPEG")
                
                okuyucu = easyocr.Reader(['ar', 'tr'], gpu=False)
                sonuc = okuyucu.readtext(gecici_yol)
                ocr_metni = " ".join([item for item in sonuc]) if sonuc else ""
                
                model = genai.GenerativeModel('gemini-1.5-flash')
                
                prompt = f"""Sen Osmanlı dönemi arşivleri ve paleografi alanında uzman kıdemli bir bilgi bilimcisin. 
                Sana sunulan bu tarihi el yazması veya matbu belgeyi analiz et ve şu protokolleri yerine getir:
                
                1. Metnin orijinal matbu/Arap harfli transkriptini eksiksiz çıkar.
                2. Metnin temiz, Latin harfli transkripsiyonlu (okunuş) halini yaz.
                3. Belge istihbaratını ve analiz verilerini tamamen Türkçe olarak ve KESİNLİKLE sadece şu JSON şablonu formatında döndür. Markdown etiketleri ekleme, doğrudan ham JSON olsun:
                   {{
                     "belge_turu": "Belge tipi (Örn: Ferman, Berat, Hüküm, Mektup, Arzuhal)",
                     "tarih_hicri": "Metinde geçen Hicri veya Rumi tarih",
                     "tarih_miladi": "Dönüştürülmüş Miladi takvim karşılığı",
                     "sahislar": ["Metinde adı veya unvanı geçen önemli tarihi kişiler"],
                     "yerler": ["Metinde adı geçen coğrafi konumlar veya bölgeler"],
                     "koordinatlar": [
                        {{"yer_adi": "İstanbul", "lat": 41.0082, "lon": 28.9784}}
                     ],
                     "ozet": "Belgenin ana konusunu anlatan tek cümlelik Türkçe özet",
                     "sozluk": {{
                        "agir_kelime_1": "günümüz_turkcesi_anlami"
                     }}
                   }}
                ÖNEMLİ: 'koordinatlar' kısmına metinde adı geçen her yerin günümüz dünya coğrafyasındaki tahmini enlem (lat) ve boylam (lon) sayısal değerlerini doğruca ekle. Yer yoksa boş bırak.
                OCR motorundan gelen kaba kelime ipuçları: '{ocr_metni}'. Maksimum doğruluğa ulaşmak için görsel bağlam ile bu ipuçlarını harmanla."""
                
                response = model.generate_content([prompt, resim_islem])
                tam_yanit = response.text
                
                metin_kismi = tam_yanit
                analiz_verisi = None
                
                if "{" in tam_yanit and "}" in tam_yanit:
                    try:
                        baslangic = tam_yanit.find("{")
                        bitis = tam_yanit.rfind("}") + 1
                        json_kismi = tam_yanit[baslangic:bitis]
                        metin_kismi = tam_yanit[:baslangic].strip()
                        analiz_verisi = json.loads(json_kismi)
                    except:
                        metin_kismi = tam_yanit
                        analiz_verisi = None
                
                sayac = len(st.session_state.belge_arsivi) + 1
                tip = "Belge"
                if analiz_verisi and isinstance(analiz_verisi, dict):
                    tip = analiz_verisi.get("belge_turu", "Belge")
                
                yeni_isim = f"📁 #{sayac} - {tip}"
                
                st.session_state.belge_arsivi[yeni_isim] = {
                    "resim": resim_islem,
                    "metin": metin_kismi if metin_kismi else tam_yanit,
                    "analiz": analiz_verisi
                }
                st.session_state.aktif_belge_adi = yeni_isim
                
                st.success("🎉 Belge analizi tamamlandı!")
                if os.path.exists(gecici_yol): os.remove(gecici_yol)
                st.rerun() 
            except Exception as e:
                st.error(f"❌ Sistem Hatası: {e}")

# ==============================================================================
# 👁️ 7. AKTİF SEÇİLİ BELGENİN EKRANA YANSITILMASI
# ==============================================================================
if st.session_state.aktif_belge_adi and st.session_state.aktif_belge_adi in st.session_state.belge_arsivi:
    
    gecerli_belge = st.session_state.belge_arsivi[st.session_state.aktif_belge_adi]
    aktif_resim = gecerli_belge["resim"]
    aktif_metin = gecerli_belge["metin"]
    aktif_analiz = gecerli_belge["analiz"]

    st.markdown(f'<div class="adim-karti">👁️ <b>ADIM 2: Çalışma Alanı — {st.session_state.aktif_belge_adi} Detayları</b></div>', unsafe_allow_html=True)
    
    col_img, col_txt = st.columns(2)
    
    with col_img:
        st.subheader("🔍 İyileştirilmiş Arşiv Görseli")
        st.image(aktif_resim, use_container_width=True)
            
    with col_txt:
        st.subheader("✍️ Düzenlenebilir Transkript Terminali")
        yeni_duzenleme = st.text_area("Metin Düzenleme Alanı", value=aktif_metin, height=380, key=f"txt_{st.session_state.aktif_belge_adi}")
        st.session_state.belge_arsivi[st.session_state.aktif_belge_adi]["metin"] = yeni_duzenleme

    if aktif_analiz and isinstance(aktif_analiz, dict):
        st.markdown('<div class="adim-karti">📊 <b>ADIM 3: Tarihsel Veri ve Coğrafi Harita Analizi</b></div>', unsafe_allow_html=True)
        
        c1, c2, c3 = st.columns(3)
        c1.metric("📜 Belge Sınıflandırması", aktif_analiz.get("belge_turu", "Bilinmiyor"))
        c2.metric("📅 Hicri Takvim", aktif_analiz.get("tarih_hicri", "Belirtilmemiş"))
        c3.metric("🌍 Miladi Takvim Karşılığı", aktif_analiz.get("tarih_miladi", "Hesaplanamadı"))
        
        st.write("---")
        st.markdown(f"💡 **Yönetici Özeti:** *{aktif_analiz.get('ozet', 'Özet mevcut değil.')}*")
        st.write("")
        
        col_metadata, col_harita = st.columns(2)
        
        with col_metadata:
            st.write("👥 **Belirlenen Tarihi Kişiler / Unvanlar:**")
            st.write(", ".join(aktif_analiz.get("sahislar", [])) if aktif_analiz.get("sahislar") else "Kişi adı ayıklanamadı.")
            st.write("")
            st.write("📍 **Ayıklanan Coğrafi Konumlar:**")
            st.write(", ".join(aktif_analiz.get("yerler", [])) if aktif_analiz.get("yerler") else "Konum bilgisi ayıklanamadı.")
            
        with col_harita:
            st.write("🗺️ **İnteraktif Belge Haritası**")
            koord_listesi = aktif_analiz.get("koordinatlar", [])
            
            if koord_listesi and isinstance(koord_listesi, list):
                try:
                    harita_df = pd.DataFrame(koord_listesi)
                    if "lat" in harita_df.columns and "lon" in harita_df.columns:
                        harita_df = harita_df.rename(columns={"lat": "latitude", "lon": "longitude"})
                        st.map(harita_df, zoom=4)
                    else:
                        st.info("Koordinat yapısı uygun bulunamadı.")
                except:
                    st.info("Harita verileri yüklenirken küçük bir pürüz oluştu.")
            else:
                st.info("Bu belgede haritalandırılacak coğrafi bir konum bulunamadı.")

        st.markdown('<div class="adim-karti">📖 <b>ADIM 3.5: Yapay Zekâ Paleografi Sözlüğü (Akıllı Lügat)</b><br>Belge metninde geçen ağır, arkaik terimlerin günümüz Türkçesi karşılıkları:</div>', unsafe_allow_html=True)
        
        aktif_sozluk = aktif_analiz.get("sozluk", {})
        if aktif_sozluk and isinstance(aktif_sozluk, dict):
            col_soz1, col_soz2 = st.columns(2)
            for i, (osmanlica_kelime, modern_anlam) in enumerate(aktif_sozluk.items()):
                hedef_kolon = col_soz1 if i % 2 == 0 else col_soz2
                with hedef_kolon:
                    st.markdown(f"""
                    <div class="sozluk-kart">
                        <b>🔑 {osmanlica_kelime.upper()} :</b> {modern_anlam}
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("Bu belgede sözlüğe eklenecek ağır veya yabancı bir terim tespit edilemedi.")

        st.markdown('<div class="adim-karti">💾 <b>ADIM 4: Dışa Aktarım ve Sertifikalı Çıktı Motoru</b><br>Raporlarınızı kurumsal formatlarda bilgisayarınıza kaydedin.</div>', unsafe_allow_html=True)
        col_word, col_pdf = st.columns(2)
        
        # Word
        doc = Document()
        doc.add_heading("PalaeoLab AI - Arşiv Analiz Raporu", 0)
        doc.add_heading("1. Çözümlenen Metin Çıktısı", level=1)
        doc.add_paragraph(yeni_duzenleme)
        
        word_akisi = BytesIO()
        doc.save(word_akisi)
        word_akisi.seek(0)
        
        with col_word:
            st.download_button(
                label="📥 Raporu Word Olarak İndir (.docx)",
                data=word_akisi,
                file_name=f"palaeolab_{st.session_state.aktif_belge_adi}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"w_{st.session_state.aktif_belge_adi}"
            )
            
        # PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="PalaeoLab AI - Onayli Veri Raporu", ln=1, align="C")
        
        temiz_metin = yeni_duzenleme.encode('latin-1', 'ignore').decode('latin-1')
        pdf.multi_cell(0, 10, txt=temiz_metin)
        pdf_akisi = pdf.output(dest='S').encode('latin-1')
        
        with col_pdf:
            st.download_button(
                label="📥 Raporu PDF Olarak İndir (.pdf)",
                data=pdf_akisi,
                file_name=f"palaeolab_{st.session_state.aktif_belge_adi}.pdf",
                mime="application/pdf",
                key=f"p_{st.session_state.aktif_belge_adi}"
            )

# ==============================================================================
# 🎛️ 5. ADIM 1: GÖRSEL İYİLEŞTİRME VE DOSYA YÜKLEME PANELİ (DEVAMI)
# ==============================================================================
    filtre_modu = st.radio("İşlem Modu", ["Renkli (Orijinal)", "Siyah-Beyaz (Yüksek Kontrast)", "Gri Tonlama"], horizontal=True)
    kontrast_oran = st.slider("Kontrast Seviyesi", 0.5, 3.0, 1.5, 0.1)
    parlaklik_oran = st.slider("Parlaklık Seviyesi", 0.5, 2.0, 1.0, 0.1)

with col_yukle:
    st.write("📂 **Belge Yükleme**")
    yuklenen_dosya = st.file_uploader("Osmanlıca belge görselini yükleyin (PNG, JPG, JPEG)", type=["png", "jpg", "jpeg"])

# ⚙️ GÖRSEL ÖN İŞLEME FONKSİYONU
def gorsel_iyilestir(image, mod, kontrast, parlaklik):
    img = image.convert("RGB")
    if mod == "Siyah-Beyaz (Yüksek Kontrast)":
        img = img.convert("L").point(lambda x: 0 if x < 128 else 255, '1')
        img = img.convert("RGB")
    elif mod == "Gri Tonlama":
        img = img.convert("L").convert("RGB")
    
    img = ImageEnhance.Contrast(img).enhance(kontrast)
    img = ImageEnhance.Brightness(img).enhance(parlaklik)
    return img

# 🔮 FONKSİYONLAR: DIŞA AKTARIM VE ANALİZ
def docx_uret(metin):
    doc = Document()
    doc.add_heading('PalaeoLab AI - Analiz Raporu', 0)
    doc.add_paragraph(metin)
    b_io = BytesIO()
    doc.save(b_io)
    return b_io.getvalue()

def pdf_uret(metin):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # UTF-8 karakter uyumluluğu için basit temizlik
    temiz_metin = metin.encode('latin-1', 'ignore').decode('latin-1')
    pdf.multi_cell(0, 10, temiz_metin)
    return pdf.output(dest='S').encode('latin-1')

@st.cache_resource
def ocr_model_yukle():
    return easyocr.Reader(['tr', 'ar'])

# 🚀 BELGE İŞLEME ANA DÖNGÜSÜ
islenmis_gorsel = None
if yuklenen_dosya is not None:
    orijinal_gorsel = Image.open(yuklenen_dosya)
    islenmis_gorsel = gorsel_iyilestir(orijinal_gorsel, filtre_modu, kontrast_oran, parlaklik_oran)
    
    col_orj, col_isl = st.columns(2)
    with col_orj:
        st.image(orijinal_gorsel, caption="Orijinal Belge", use_container_width=True)
    with col_isl:
        st.image(islenmis_gorsel, caption="İyileştirilmiş Belge", use_container_width=True)
        
    if st.button("🔮 Belgeyi Arşive Ekle ve Analize Başla", type="primary"):
        b_name = yuklenen_dosya.name
        if b_name not in st.session_state.belge_arsivi:
            st.session_state.belge_arsivi[b_name] = {"gorsel": islenmis_gorsel, "analiz": None}
        st.session_state.aktif_belge_adi = b_name
        st.rerun()

# ==============================================================================
# 🧠 6. ADIM 2: YAPAY ZEKA VE PALEOGRAFİ ANALİZ LABORATUVARI
# ==============================================================================
if st.session_state.aktif_belge_adi and st.session_state.aktif_belge_adi in st.session_state.belge_arsivi:
    aktif_veri = st.session_state.belge_arsivi[st.session_state.aktif_belge_adi]
    
    st.markdown(f"""
    <div class="adim-karti">
        🧠 <b>ADIM 2: Yapay Zekâ Paleografi Analiz Odası</b><br>
        Şu an analiz edilen belge: <code>{st.session_state.aktif_belge_adi}</code>
    </div>
    """, unsafe_allow_html=True)
    
    col_buton1, col_buton2 = st.columns(2)
    
    with col_buton1:
        if st.button("👁️ EasyOCR ile Ön Karakter Taraması Yap"):
            with st.spinner("Görsel üzerindeki metin katmanları taranıyor..."):
                reader = ocr_model_yukle()
                islenmis_gorsel.save("gecici.png")
                sonuc = reader.readtext("gecici.png", detail=0)
                if os.path.exists("gecici.png"): os.remove("gecici.png")
                
                ocr_metni = " ".join(sonuc)
                st.session_state.belge_arsivi[st.session_state.aktif_belge_adi]["ocr_ham"] = ocr_metni
                st.success("Ön tarama tamamlandı! Sayfa altına eklenen alandan ham veriyi görebilirsiniz.")

    with col_buton2:
        if st.button("🤖 Gemini AI ile Derin Paleografik Analiz Başlat"):
            if "GEMINI_API_KEY" in st.secrets:
                with st.spinner("Gemini multimodal katmanı belgeyi çözümlüyor..."):
                    try:
                        model = genai.GenerativeModel('gemini-2.5-flash')
                        
                        # Görsel verisini Gemini formatına dönüştürme
                        b_io = BytesIO()
                        aktif_veri["gorsel"].save(b_io, format="PNG")
                        gorsel_parca = {"mime_type": "image/png", "data": b_io.getvalue()}
                        
                        prompt = """
                        Sen uzman bir Osmanlı paleografi uzmanı ve tarihçisin. Ekli Osmanlıca belge görselini incele ve şu adımları eksiksiz gerçekleştir:
                        1. Transkripsiyon: Metnin Osmanlıca okunuşunu (Latin harfleriyle) satır satır çıkar.
                        2. Günümüz Türkçesi: Metni günümüz akıcı Türkçesine sadeleştirerek çevir.
                        3. Tarihsel Analiz: Belgenin türünü (Ferman, Berat, Ariza vb.), dönemini, varsa tarihi ve kurumları analiz et.
                        4. Mini Sözlük: Belgede geçen ağır Arapça, Farsça veya eski Türkçe terimlerden en az 5 tanesini seçip anlamlarını yaz.
                        Yanıtı profesyonel Markdown formatında, başlıklar kullanarak ver.
                        """
                        
                        yanit = model.generate_content([prompt, gorsel_parca])
                        st.session_state.belge_arsivi[st.session_state.aktif_belge_adi]["analiz"] = yanit.text
                        st.rerun()
                    except Exception as e:
                        st.error(f"Gemini API hatası oluştu: {str(e)}")
            else:
                st.error("Lütfen önce Streamlit Secrets alanına geçerli bir GEMINI_API_KEY tanımlayın.")

    # 📊 SONUÇLARIN GÖSTERİLMESİ
    if "ocr_ham" in aktif_veri and aktif_veri["ocr_ham"]:
        with st.expander("📝 EasyOCR Tarafından Yakalanan Ham Metin Katmanı"):
            st.code(aktif_veri["ocr_ham"], language="text")

    if aktif_veri["analiz"]:
        st.markdown("### 📊 Yapay Zekâ Analiz ve Deşifre Raporu")
        
        # Sonuç ekranını düzenlenebilir kılmak için TextArea'ya gömme seçeneği
        rapor_metni = st.text_area("Düzenlenebilir Rapor Çıktısı", value=aktif_veri["analiz"], height=450)
        st.session_state.belge_arsivi[st.session_state.aktif_belge_adi]["analiz"] = rapor_metni
        
        st.write("---")
        st.markdown("### 📥 Raporu Dışa Aktar")
        col_down1, col_down2 = st.columns(2)
        
        with col_down1:
            docx_data = docx_uret(rapor_metni)
            st.download_button(
                label="📥 Word Dosyası (.docx) Olarak İndir",
                data=docx_data,
                file_name=f"PalaeoLab_{st.session_state.aktif_belge_adi}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        with col_down2:
            pdf_data = pdf_uret(rapor_metni)
            st.download_button(
                label="📥 PDF Dosyası (.pdf) Olarak İndir",
                data=pdf_data,
                file_name=f"PalaeoLab_{st.session_state.aktif_belge_adi}.pdf",
                mime="application/pdf",
                use_container_width=True
            )

# ==============================================================================
# 🔮 GELİŞTİRİLMİŞ DIŞA AKTARIM MOTORLARI VE PREMİUM PROMPT KATMANI
# ==============================================================================

def docx_uret(metin):
    """Gelişmiş Word Rapor Üreticisi"""
    doc = Document()
    # Başlık stili
    baslik = doc.add_heading('🔬 PalaeoLab AI - Paleografik Analiz ve Deşifre Raporu', 0)
    baslik.alignment = 1 # Ortalanmış
    
    doc.add_paragraph("Bu rapor PalaeoLab AI paleografi otomasyon sistemi tarafından üretilmiştir.")
    doc.add_paragraph("-" * 40)
    
    # Rapor içeriğini ekle
    doc.add_paragraph(metin)
    
    b_io = BytesIO()
    doc.save(b_io)
    return b_io.getvalue()

def pdf_uret(metin):
    """Türkçe Karakter Uyumlu PDF Üreticisi"""
    pdf = FPDF()
    pdf.add_page()
    
    # Standart fontlar yerine evrensel Unicode/Latin1 güvenli karakter temizliği
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(0, 10, "PalaeoLab AI - Analiz Raporu", ln=1, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 5, "Tarih: 2026-05-19 | Sistem: Evrensel Arsiv ve Analiz Katmani", ln=1, align="C")
    pdf.ln(10)
    
    pdf.set_font("Helvetica", size=11)
    
    # Türkçe karakterlerin PDF çıktısında kırılmaması için güvenli harf dönüşüm haritası
    turkce_harfler = {
        'ı': 'i', 'ğ': 'g', 'ü': 'u', 'ş': 's', 'ö': 'o', 'ç': 'c',
        'İ': 'I', 'Ğ': 'G', 'Ü': 'U', 'Ş': 'S', 'Ö': 'O', 'Ç': 'C'
    }
    
    temiz_metin = metin
    for kaynak, hedef in turkce_harfler.items():
        temiz_metin = temiz_metin.replace(kaynak, hedef)
        
    # Satır satır yazdırarak sayfa taşma kontrolü sağlama
    for satir in temiz_metin.split('\n'):
        if satir.strip() == "":
            pdf.ln(4)
        else:
            pdf.multi_cell(0, 7, satir)
            
    return pdf.output(dest='S').encode('latin-1', 'ignore')

# 🏛️ AKADEMİK GEMİNİ PALEOGRAFİ PROMPTU
PALEOGRAFI_PROMPTU = """
Sen T.C. Cumhurbaşkanlığı Devlet Arşivleri standartlarında çalışan kıdemli bir Osmanlı Paleografisi uzmanı, diplomatik tarihçi ve epigraftsın. 
Görseldeki Osmanlı Türkçesi (Arabi harfli) belgeyi diplomatik kurallara göre deşifre et ve yapılandırılmış bir akademik rapor hazırla.

Yanıtını tam olarak şu başlıklarla ve Markdown formatında sun:

### 📑 1. DİPLOMATİK VE TÜR ANALİZİ
*   **Belge Türü:** (Ferman, Berat, Ariza, Telgraf, Hüccet, İrade-i Seniyye, tahrirat vb. hangisi olduğunu gerekçesiyle yaz.)
*   **Yazı Türü (Hat):** (Rika, Divani, Nesih, Talik, Sülüs, Siyakat vb. tahmin et.)
*   **Kurumsal Aidiyet:** (Belgenin çıktığı ve gittiği devlet kurumları, makamlar.)

### 📝 2. TRANSKRİPSİYON (METNİN OKUNUŞU)
*Belgenin orijinal Osmanlıca okunuşunu (Latin harfleriyle transkripsiyon kurallarına uygun olarak) satır satır veya düzenli paragraflar halinde buraya aktar. Bilinmeyen/okunamayan kelimeler için [okunamadı] ifadesini kullan.*

### 🔄 3. GÜNÜMÜZ TÜRKÇESİNE SADELEŞTİRME
*Metnin içeriğini, tarihsel bağlamını ve hukuki/idari anlamını bozmadan, günümüz resmi ve akıcı Türkçesiyle tam metin olarak özetle/çevir.*

### ⏳ 4. TARİH VE TAKVİM DÖNÜŞÜMÜ
*Belgede geçen Hicri veya Rumi tarihleri tespit et. Bu tarihleri Miladi takvime (gün-ay-yıl netliğinde) dönüştür.*

### 📖 5. ARŞİV VE TERİMLER SÖZLÜĞÜ
*Belgede geçen unvanlar, devlet görevleri, hukuki terimler veya ağır Arapça/Farsça tamlamalardan en az 5 tanesini seçerek anlamlarını açıkla.*
(Örnek: "Bende-i dâ'î", "Mîr-i mîrân", "Mûcebince amel oluna" vb.)
"""
